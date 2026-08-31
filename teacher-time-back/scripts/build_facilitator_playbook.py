from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "Teacher-Time-Back-Lab-Facilitator-Playbook.docx"

DEEP = "063C33"
TEAL = "008C84"
AQUA = "59DDD0"
GOLD = "F2B633"
PAPER = "F4EFE4"
MINT = "E8F6F3"
INK = "123A35"
MUTED = "4F6B67"
WHITE = "FFFFFF"
GRAY = "E5ECE9"
CAUTION = "FFF4D8"
RISK = "FDEBE7"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text, url, color=TEAL):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    keep.set(qn("w:val"), "1" if value else "0")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])
    set_run_font(run, size=8.5, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.17

    for style_name, size, color, before, after in (
        ("Title", 30, DEEP, 0, 6),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 19, DEEP, 16, 8),
        ("Heading 2", 14.5, TEAL, 13, 6),
        ("Heading 3", 11.5, DEEP, 9, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial Narrow" if style_name != "Subtitle" else "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name not in {"Subtitle"}
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.17

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("TEACHER TIME-BACK LAB  |  FACILITATOR PLAYBOOK")
    set_run_font(run, "Arial Narrow", 8.5, TEAL, True)
    p.paragraph_format.space_after = Pt(0)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.86))
    set_table_geometry(table, [7000, 2880], indent_dxa=0)
    left = table.cell(0, 0).paragraphs[0]
    left_run = left.add_run("Andrew Wolverton  |  Human-led AI workflow design")
    set_run_font(left_run, size=8.5, color=MUTED)
    add_page_number(table.cell(0, 1).paragraphs[0])


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_run_font(r, "Arial Narrow", 10, TEAL, True)
    return p


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph(style="Title")
    p.add_run(text)
    if subtitle:
        p2 = doc.add_paragraph(style="Subtitle")
        p2.add_run(subtitle)


def add_callout(doc, label, text, fill=MINT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9880], indent_dxa=0)
    cell = table.cell(0, 0)
    set_row_cant_split(table.rows[0])
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    set_run_font(r, "Arial Narrow", 9.5, accent, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    body_color = WHITE if fill == DEEP else INK
    set_run_font(r2, "Arial", 10.5, body_color, True if label.lower() in {"the rule", "decision"} else False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, bold_start=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label.upper() + "  ")
    set_run_font(r, "Arial Narrow", 9.5, TEAL, True)
    p.add_run(text)
    return p


def add_script_block(doc, say, show=None, do=None, listen=None, decision=None):
    table = doc.add_table(rows=0, cols=2)
    entries = [("SAY", say)]
    if show:
        entries.append(("SHOW", show))
    if do:
        entries.append(("PARTICIPANTS", do))
    if listen:
        entries.append(("LISTEN FOR", listen))
    if decision:
        entries.append(("DECISION", decision))
    for label, text in entries:
        cells = table.add_row().cells
        set_cell_fill(cells[0], DEEP)
        set_cell_fill(cells[1], "F8FAF9")
        p1 = cells[0].paragraphs[0]
        r1 = p1.add_run(label)
        set_run_font(r1, "Arial Narrow", 9, WHITE, True)
        p2 = cells[1].paragraphs[0]
        r2 = p2.add_run(text)
        set_run_font(r2, "Arial", 10.2, INK)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, [1500, 8380], indent_dxa=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_faq(doc, question, short_answer, explanation, route, sources=None):
    q = doc.add_paragraph(style="Heading 3")
    q.add_run(question)
    p = doc.add_paragraph()
    r = p.add_run("Short answer: ")
    set_run_font(r, bold=True, color=DEEP)
    p.add_run(short_answer)
    p2 = doc.add_paragraph(explanation)
    add_label_paragraph(doc, "Route", route)
    if sources:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(5)
        r3 = p3.add_run("Official guidance: ")
        set_run_font(r3, size=8.5, color=MUTED, bold=True)
        for idx, (label, url) in enumerate(sources):
            if idx:
                p3.add_run(" | ")
            add_hyperlink(p3, label, url)


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
style_document(doc)

# Cover
add_kicker(doc, "Facilitator playbook")
add_title(
    doc,
    "Teacher Time-Back Lab",
    "Complete run of show, speaker script, safety decisions, nuanced FAQ, and follow-up pathways",
)
add_callout(
    doc,
    "The promise",
    "Teachers do not come to master AI. They come to stop rebuilding one recurring piece of work, safely and with professional judgment intact.",
    fill=DEEP,
    accent=AQUA,
)

metric = doc.add_table(rows=1, cols=4)
set_table_geometry(metric, [2470, 2470, 2470, 2470], indent_dxa=0)
for idx, (big, small) in enumerate((("45 min", "Leadership setup"), ("90 min", "Teacher build lab"), ("2 weeks", "Measured use"), ("30 min", "Follow-up decision"))):
    cell = metric.cell(0, idx)
    set_cell_fill(cell, CAUTION if idx % 2 == 0 else MINT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(big)
    set_run_font(r, "Arial Narrow", 15, DEEP, True)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(small)
    set_run_font(r2, size=8.5, color=MUTED)

doc.add_paragraph()
add_label_paragraph(doc, "Audience", "Andrew Wolverton, school hosts, instructional technology leaders, and approved co-facilitators.")
add_label_paragraph(doc, "Current unit boundary", "Beginner, student-neutral assistant design. No student records are needed or used.")
add_label_paragraph(doc, "Important", "This playbook is operational guidance, not legal advice or a compliance certification. School policy and the responsible privacy, data, legal, special-education, and technology leaders control local decisions.")

add_page_break(doc)

# Navigation
add_kicker(doc, "How to use this playbook")
add_title(doc, "Follow the room, not the page count.", "Each stage tells you what to say, show, ask participants to do, listen for, and decide.")
for text in (
    "Part 1. Offer and outcome: what is actually being sold and built.",
    "Part 2. Leadership setup: the decisions that must exist before teachers arrive.",
    "Part 3. The 90-minute lab: a complete minute-by-minute facilitation script.",
    "Part 4. Safety rehearsal: four interactive choices and their teaching points.",
    "Part 5. Follow-up: how teachers decide to keep, revise, or stop.",
    "Part 6. Nuanced FAQ: student work, coded records, groups, trends, assessment data, and approved tools.",
    "Part 7. Advanced PD pathway: where student-level and group-level analysis could responsibly go next.",
    "Appendix: quick-reference decision ladder, facilitator language, and official sources.",
):
    add_bullet(doc, text)

add_callout(doc, "Facilitation rule", "When a participant asks whether a student-data workflow is allowed, do not improvise a yes or no. Name the risk category, explain the decision gates, and route the local permission question to the school owner.")

# Part 1
add_page_break(doc)
doc.add_heading("Part 1 | What this lab actually delivers", level=1)
doc.add_paragraph("This is a bounded professional learning product, not an AI feature tour. The teacher leaves with a reusable assistant blueprint for one recurring task. The assistant drafts or organizes. The educator makes every instructional, evaluative, communication, and final-use decision.")

doc.add_heading("The finished teacher product", level=2)
for text in (
    "One job: a single repeated task stated in one sentence.",
    "Allowed sources: blank, public, fictional, educator-owned, or school-approved student-neutral material.",
    "Exact output: a predictable structure the teacher can review quickly.",
    "Quality criteria: what makes the draft useful in this teacher's context.",
    "Refusal behavior: when the assistant must stop and request a safe replacement.",
    "Human decision line: what remains exclusively with the teacher.",
    "Fictional test: a rehearsal that does not require student records.",
    "Measurement plan: total task time, correction time, usefulness, safety slips, and keep/revise/stop.",
):
    add_bullet(doc, text)

doc.add_heading("What teachers need before they arrive", level=2)
for text in (
    "One recurring task they personally perform.",
    "A laptop or a printed workbook.",
    "Access to the exact school-approved AI environment, if leadership has named one.",
    "Professional judgment and willingness to review a draft.",
):
    add_bullet(doc, text)

doc.add_heading("What teachers do not need", level=2)
for text in (
    "AI vocabulary, coding, prompt-engineering experience, or a prebuilt assistant.",
    "Student names, records, identifiable work, grades, attendance, behavior, IEP, 504, counseling, health, discipline, family, personnel, or other confidential data.",
    "A personal account or a workaround when the approved environment is unclear.",
):
    add_bullet(doc, text)

add_callout(doc, "Success measure", "The goal is not maximum AI use. The goal is one safe, understandable workflow that measurably reduces total work after review and correction.")

# Part 2
add_page_break(doc)
doc.add_heading("Part 2 | Leadership setup: 45 minutes", level=1)
add_callout(doc, "Required before the lab", "If the school cannot name the exact approved AI product, account type, access route, and local escalation contact, the session runs in paper/demo mode.")

segments = [
    ("0:00-0:05 | Establish the decision", "This is not a policy-writing meeting and not a vendor demo. We are creating a safe container for one beginner workshop. Teachers need the exact approved environment, visible input boundaries, a human decision line, and a named person for questions outside the lab.", "Confirm sponsor, cohort, date, format, and local decision owner."),
    ("0:05-0:15 | Name the approved environment", "What exact product may teachers use? What account or tenant? How do they verify that they are in it? What features or data categories are excluded? Who owns an uncertain answer?", "Write the exact tool and account route on the opening slide. If unknown, select paper/demo mode."),
    ("0:15-0:25 | Set the input lanes", "Tool approval and task approval are different questions. An approved account does not make every input or every use appropriate.", "Confirm local examples for green, pause-and-ask, and never-enter categories."),
    ("0:25-0:35 | Draw the human decision line", "The assistant may draft, organize, reformat, summarize approved neutral material, or propose options. It does not make student-level or regulated decisions.", "Name the required accuracy, bias, appropriateness, and alignment review."),
    ("0:35-0:42 | Choose the first-build lanes", "Which two or three lower-risk tasks match this cohort and can be verified quickly without student records?", "Select Reuse Planner, neutral Family Message Drafter, approved Meeting-to-Action Assistant, Teaching Library Organizer, or Feedback Framework Builder."),
    ("0:42-0:45 | Confirm day-of ownership", "Who opens the session, handles sign-in, answers local questions, captures unresolved issues, and schedules the follow-up?", "Complete the readiness checklist before promoting the lab."),
]
for heading, say, decision in segments:
    doc.add_heading(heading, level=2)
    add_script_block(doc, say=say, decision=decision)

doc.add_heading("Leadership readiness checklist", level=2)
for text in (
    "Exact approved product and account named.",
    "Data and task restrictions available in plain language.",
    "Privacy/data owner and technology support contact named.",
    "School host knows the lab uses no student records.",
    "Fictional practice pack and paper/demo fallback ready.",
    "Two-week follow-up date reserved.",
):
    add_bullet(doc, "[ ] " + text)

# Part 3
add_page_break(doc)
doc.add_heading("Part 3 | Teacher build lab: 90 minutes", level=1)

lab_steps = [
    {
        "heading": "0:00-0:05 | Welcome and lower the entry point",
        "say": "You are not here to master AI. You are here to stop rebuilding one piece of work. You need one repeated task, the school's approved environment, and your professional judgment. I will give you the structure. We will not use student records today.",
        "show": "The opening promise and one finished Reuse Planner example.",
        "do": "Complete the sentence: I keep rebuilding, rewriting, reorganizing, or carrying home...",
        "listen": "Broad tasks such as planning, emails, grading, meetings. Do not correct yet. Help each person name one recurring burden.",
        "decision": "Everyone has a time drain to examine.",
    },
    {
        "heading": "0:05-0:12 | Show the finished product first",
        "say": "This is what you are leaving with. It is more durable than a clever one-time prompt because it names the job, sources, output, stop line, and the decision only you can make.",
        "show": "The six fields in the Completed Reuse Planner Example and one fictional input/output pair.",
        "do": "Circle the field that would make an AI draft easiest for you to check.",
        "listen": "Questions about whether this is a prompt, an assistant, or a tool. Answer: it is a reusable instruction blueprint implemented only in the approved environment.",
        "decision": "Participants can describe the product in plain language.",
    },
    {
        "heading": "0:12-0:25 | Step 1: Choose one task",
        "say": "The first assistant should be boring in the best possible way. Choose a repeated, bounded task you can check quickly.",
        "show": "The four filters: repeatable, student-neutral, reviewable, bounded.",
        "do": "List up to three time drains. Circle one that passes all four filters. Estimate frequency, current minutes, and what the recovered time would be for.",
        "listen": "Tasks that quietly depend on student records, make student decisions, or are too broad. Narrow or reroute them.",
        "decision": "Each participant has one job stated in one sentence.",
    },
    {
        "heading": "0:25-0:40 | Step 2: Protect the inputs",
        "say": "Before we write instructions, we decide what is allowed through the door. A name is not the only identifier. Details can identify someone alone or in combination.",
        "show": "Green, pause-and-ask, and never-enter cards plus the exact approved product/account.",
        "do": "List allowed sources, never-enter categories, and the person to ask when unsure.",
        "listen": "Names removed but rare details retained, personal accounts, real student work, or uncertainty about tool approval.",
        "decision": "Every participant has both an allowed-input list and a never-enter list. Unknown approval means blueprint mode only.",
    },
    {
        "heading": "0:40-0:65 | Step 3: Build the blueprint",
        "say": "We are not chasing a magic phrase. We are writing a visible contract for one job.",
        "show": "Role, sources, output, quality, refusal, and human decision.",
        "do": "Complete all six fields and assemble the copy-ready instruction block.",
        "listen": "Vague roles, unbounded sources, hidden student-data needs, missing refusal behavior, and unclear teacher ownership.",
        "decision": "Each participant has one complete, copy-ready instruction block.",
    },
    {
        "heading": "0:65-0:80 | Step 4: Practice safely",
        "say": "We test the structure before we trust the workflow. Today's content is fictional on purpose.",
        "show": "The Harbor City Fictional Practice Pack and the planted unsupported bus-stop claim.",
        "do": "Run or role-play the test. Compare the output to the source. Find the invented claim. Review accuracy, bias, appropriateness, and alignment. Revise one instruction.",
        "listen": "Participants accepting polished output, overlooking missing facts, or fixing the output without strengthening the instruction.",
        "decision": "Everyone identifies one correction and makes one blueprint revision.",
    },
    {
        "heading": "0:80-0:88 | Step 5: Measure a real result",
        "say": "Time saved is not generation time. It is the full task time after checking and correcting the output.",
        "show": "Current minutes, assisted minutes, correction minutes, usefulness, safety slips, and keep/revise/stop.",
        "do": "Design a two-week test and record a decision rule.",
        "listen": "Guaranteed time-savings expectations or using frequency of AI use as success.",
        "decision": "Each participant can explain what evidence would make them keep, revise, or stop.",
    },
    {
        "heading": "0:88-0:90 | Close",
        "say": "The assistant drafts. You decide. You now have one job, one safe source lane, one refusal line, one fictional test, one review routine, and one way to decide whether this is worth keeping.",
        "show": "The completed workbook checklist and follow-up date.",
        "do": "Save the blueprint and choose the first date to test it.",
        "decision": "No student data is collected. Only cohort-level follow-up commitments are recorded.",
    },
]

for step in lab_steps:
    doc.add_heading(step["heading"], level=2)
    add_script_block(doc, step["say"], step.get("show"), step.get("do"), step.get("listen"), step.get("decision"))

doc.add_heading("Copy-ready build language", level=2)
add_callout(doc, "Facilitator model", "You are my [assistant name]. Your only job is to [one bounded task]. Use only [approved source types]. Return [exact structure]. A useful result must [quality criteria]. If the request contains [prohibited categories] or asks for [prohibited decisions], stop without repeating the information and ask for a safe replacement. You may draft or organize. I decide [human decision]. Remind me to review accuracy, bias, appropriateness, and alignment before use.")

doc.add_heading("Facilitator working notes", level=2)
doc.add_paragraph("Keep this page open during the build. Record cohort patterns, not student information or identifiable participant examples.")
for label in (
    "Likely time-drain tasks",
    "Approved environment and local boundaries",
    "Questions to route to the school owner",
    "Two-week follow-up commitments",
):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9880], indent_dxa=0)
    set_cell_fill(table.cell(0, 0), "F8FAF9")
    paragraph = table.cell(0, 0).paragraphs[0]
    run = paragraph.add_run(label.upper())
    set_run_font(run, "Arial Narrow", 9, TEAL, True)
    table.cell(0, 0).add_paragraph("\n")
    set_row_cant_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Part 4
add_page_break(doc)
doc.add_heading("Part 4 | Interactive safety rehearsal", level=1)
doc.add_paragraph("This is a four-question guided rehearsal, not a certification. The interface shows one question at a time. The correct choice reveals the safe visual state and unlocks the next question. Wrong choices receive calm, explanatory feedback.")

questions = [
    ("1. Account check", "A teacher is ready to start. What comes first?", "Confirm the exact school-approved tool and account.", "A personal account is not a workaround when the approved environment is unknown."),
    ("2. Identifier check", "The name is removed, but rare details remain. What now?", "Stop and replace the material with fictional or student-neutral content.", "Direct identifiers are not the entire risk. Combinations may still point to a person."),
    ("3. Task check", "A teacher wants AI to grade identifiable student essays. What is the safer first build?", "Build a feedback-framework assistant using criteria and fictional examples.", "Begin with a lower-risk drafting aid, not a student-level decision."),
    ("4. Output check", "The draft looks polished. What happens before use?", "Review accuracy, bias, appropriateness, and alignment.", "Professional formatting is not evidence that an output is correct."),
]
for heading, prompt, answer, point in questions:
    doc.add_heading(heading, level=2)
    add_label_paragraph(doc, "Prompt", prompt)
    add_label_paragraph(doc, "Safer move", answer)
    add_label_paragraph(doc, "Teaching point", point)

# Part 5
add_page_break(doc)
doc.add_heading("Part 5 | Two-week follow-up: 30 minutes", level=1)
follow = [
    ("0:00-0:05 | Reopen the decision", "This is not a celebration of usage. We are deciding what was genuinely useful, what needs revision, and what should stop."),
    ("0:05-0:15 | Compare the full task time", "Compare current minutes, assisted minutes, and correction minutes. Ask what teachers did with any time recovered."),
    ("0:15-0:23 | Diagnose patterns", "Sort failures into unclear task, unsafe or vague sources, weak output format, invented/missing content, too much correction, unclear policy, or wrong task."),
    ("0:23-0:28 | Decide", "Each teacher chooses keep, revise, or stop and records one evidence-based reason."),
    ("0:28-0:30 | Route unresolved needs", "Leadership records cohort-level support and policy questions. Results are not used for teacher evaluation."),
]
for heading, say in follow:
    doc.add_heading(heading, level=2)
    add_script_block(doc, say=say)

# Part 6 FAQ
add_page_break(doc)
doc.add_heading("Part 6 | Nuanced questions and facilitator answers", level=1)
add_callout(doc, "The rule", "A teacher may have a legitimate educational reason to review student information. That does not automatically authorize disclosure of the information to an AI vendor or use in a particular AI feature. Tool, account, data, purpose, access, contract, and local policy must all align.", fill=CAUTION, accent=GOLD)
doc.add_heading("Frequently asked questions", level=2)

FERPA_PII = ("FERPA PII definition", "https://studentprivacy.ed.gov/content/personally-identifiable-information-education-records")
FERPA_DEID = ("FERPA de-identification FAQ", "https://studentprivacy.ed.gov/faq/what-constitutes-de-identified-records-and-information")
FERPA_RULE = ("FERPA regulations", "https://studentprivacy.ed.gov/ferpa")
NYSED = ("NYSED Education Law 2-d", "https://www.nysed.gov/data-privacy-security/parents-and-students")
NYCPS = ("NYCPS AI guidance", "https://www.schools.nyc.gov/about-us/vision-and-mission/artificial-intelligence/guidance-on-artificial-intelligence-full")
NYCPS_ERMA = ("NYCPS ERMA overview", "https://www.schools.nyc.gov/about-us/vision-and-mission/artificial-intelligence/guidance-on-artificial-intelligence")

faq_items = [
    (
        "Can a teacher use AI to review specific student work?",
        "Not in this beginner lab. In practice, it may be possible only inside a school-approved workflow with a legitimate educational purpose and explicit permission for that tool, account, data type, and task.",
        "Student work can be an education record and can contain direct or indirect identifiers. Removing the student's name may not remove identity. The teacher's authority to review the work does not automatically extend to disclosing it to a third-party AI system.",
        "Current lab: use fictional or student-neutral examples. Advanced workflow: require the school's instructional, privacy/data, and technology owners to approve the product and procedure before any live data is used.",
        [FERPA_PII, NYCPS_ERMA],
    ),
    (
        "What if the teacher replaces the student ID with a random key?",
        "A random key can be one part of a formal de-identification design, but replacing one identifier does not automatically make the dataset de-identified or approved for AI use.",
        "If the teacher or school retains a key that reconnects the row to a student, or if the remaining details can identify a student, the data may still be linkable. FERPA permits coded de-identified student-level data for education research under specific conditions, including that the recipient cannot identify the student and does not receive the code-generation method. That rule is not a do-it-yourself permission for teachers to upload coded rows to an AI product.",
        "Treat keyed row-level data as protected until the school data/privacy owner makes and documents the de-identification and tool-use determination.",
        [FERPA_DEID, FERPA_RULE],
    ),
    (
        "Does removing names make student data anonymous?",
        "No.",
        "FERPA PII includes direct identifiers, indirect identifiers, and combinations of information that allow a reasonable person in the school community to identify a student with reasonable certainty. A rare program, date, grade, language, disability category, event, or small group can make a record recognizable.",
        "Do not teach name removal as the safety standard. Use fictional material in this lab and formal de-identification processes in advanced work.",
        [FERPA_PII, FERPA_DEID],
    ),
    (
        "Can AI help a teacher identify class-wide learning trends?",
        "Potentially, but the approved workflow matters more than the analytical goal.",
        "NYCPS guidance recognizes that AI may surface patterns in data while educators interpret findings. If the input includes student data, NYCPS also requires an ERMA-approved tool. A safer starting point is teacher-created aggregate counts that do not expose small groups or identifiable rows. Even aggregate data can reveal students when the class or subgroup is small.",
        "Possible advanced PD. Require the approved analytics/AI environment, minimum necessary data, aggregation rules, small-group safeguards, and educator interpretation.",
        [NYCPS, NYCPS_ERMA],
    ),
    (
        "Can a teacher track different groups of students with AI?",
        "This is not automatically prohibited, but it is a higher-risk workflow and should not be introduced as an individual teacher experiment.",
        "Groups may reveal protected or sensitive characteristics, and small groups can make students identifiable. The purpose also matters: noticing an instructional pattern is different from ranking, labeling, predicting, placing, disciplining, or allocating opportunities. Bias and proxy variables can distort the result.",
        "Advanced PD only after leadership defines permissible purposes, approved data fields, minimum group size or suppression rules, decision limits, review, documentation, and escalation.",
        [FERPA_PII, NYCPS],
    ),
    (
        "Can AI recommend which students belong in intervention groups?",
        "Not as an autonomous recommendation in Andrew's current training.",
        "Grouping can affect student opportunity and may rely on incomplete or biased signals. AI can help an educator organize an approved, non-identifiable set of instructional patterns, but the teacher and required school team must make grouping, intervention, placement, disability, and service decisions under local procedures.",
        "Keep the current lab focused on reusable materials. Route student grouping or intervention decisions to leadership and qualified specialists for a separate governed workflow.",
        [NYCPS],
    ),
    (
        "Can a teacher use AI for assessment item analysis?",
        "Yes as a concept, but live student response data requires a separately approved workflow.",
        "A low-risk version analyzes the assessment questions, answer key, and fictional or aggregate response counts. A higher-risk version uploads row-level responses or student work. The latter requires approval for the tool, data, purpose, and access.",
        "Current lab: use fictional counts or teacher-created non-identifiable summaries. Advanced PD: use only the school's authorized data source and approved environment.",
        [FERPA_DEID, NYCPS],
    ),
    (
        "Can AI compare a student's performance over time?",
        "That is student-level longitudinal analysis and sits outside this beginner unit.",
        "A key code does not eliminate linkability when multiple records are intentionally connected over time. Longitudinal detail can increase re-identification risk and may influence high-impact decisions.",
        "Require a school-owned analytics process, an approved product, legitimate educational interest, minimum necessary access, and professional review. Do not build this in a personal AI account.",
        [FERPA_RULE, NYCPS_ERMA],
    ),
    (
        "Can a teacher upload a spreadsheet exported from the SIS or LMS?",
        "Not unless the school has explicitly approved that product, account, feature, data type, and purpose.",
        "Exports often contain identifiers or combinations of grades, attendance, behavior, demographics, accommodations, or dates. Deleting a name column is not enough. Copying data from an approved system into another product is a new disclosure and processing context.",
        "Use school-provided reporting tools first. Route any proposed AI upload through the responsible data/privacy and technology owners.",
        [FERPA_PII, NYSED, NYCPS_ERMA],
    ),
    (
        "If a product says it is FERPA compliant, can a teacher use it?",
        "No vendor claim replaces school approval.",
        "FERPA does not operate as a universal product certification. The school's relationship with the vendor, direct control, legitimate educational purpose, contract terms, data handling, account configuration, and state/local rules matter. In NYCPS, student or staff data may not be entered into tools that have not completed the applicable ERMA review.",
        "Ask leadership to confirm the exact approved product, account, and permitted use in the official portal or local process.",
        [NYSED, NYCPS_ERMA],
    ),
    (
        "Does an approved tool make every AI use acceptable?",
        "No.",
        "Approval may cover privacy and security while leaving instructional appropriateness, bias, age restrictions, task design, professional responsibility, or prohibited decisions unresolved. NYCPS explicitly states that ERMA approval is not the only requirement for use.",
        "Check both the environment and the task. Keep the human decision line visible.",
        [NYCPS_ERMA, NYCPS],
    ),
    (
        "Can the teacher ask AI to de-identify the data first?",
        "Not by uploading identifiable data into an unapproved tool.",
        "The disclosure has already occurred when the data is entered. Effective de-identification also requires attention to direct identifiers, indirect identifiers, combinations, other available data, and repeated releases.",
        "De-identify inside the school-approved process before the AI system receives anything, with the school's responsible data/privacy owner defining the method.",
        [FERPA_DEID, FERPA_RULE],
    ),
    (
        "Can AI grade, score, or rank students if the teacher reviews the result?",
        "Not in this lab and not merely because a human signs off afterward.",
        "Human review is necessary but does not automatically make a high-impact workflow appropriate. Grading, ranking, placement, discipline, accommodations, and service decisions need stronger validation, governance, transparency, and local authority.",
        "Use AI to help build a rubric structure or neutral feedback stems from approved criteria. The teacher reads the work, evaluates it, responds, and assigns every score.",
        [NYCPS],
    ),
    (
        "Can a teacher use AI to draft feedback for a particular student?",
        "Not with identifiable student work in this beginner lab.",
        "A safer first build creates feedback frameworks, question stems, exemplars, or comment banks from approved criteria and fictional examples. Applying those materials to a student's work remains the teacher's job unless the school has approved a specific student-data workflow.",
        "Current lab: framework building only. Advanced workflow: requires approval for identifiable work and the exact feedback tool.",
        [FERPA_PII, NYCPS_ERMA],
    ),
    (
        "Can a teacher enter IEP, 504, disability, counseling, health, behavior, or discipline information?",
        "No in this lab. These requests must be routed to qualified staff and the school's approved systems and procedures.",
        "These categories are highly sensitive, may be protected by multiple laws and local rules, and may influence regulated decisions. Andrew's training does not determine accommodations, eligibility, services, placement, discipline, counseling, or safety responses.",
        "Stop, do not repeat the information, use fictional material for training, and refer the real question to the responsible specialist and privacy/data owner.",
        [FERPA_PII, NYSED],
    ),
    (
        "Can AI draft family communications?",
        "It can help with neutral, general communications in an approved environment, with human review. Student-specific communications require additional permission and care.",
        "General dates, logistics, and educator-written information can fit a lower-risk lane. Student-specific performance, behavior, disability, family, or other protected information changes the risk. Translation adds another accuracy and impact review need.",
        "Current lab: neutral logistics and general updates only. Student-specific communication belongs in a separately approved workflow.",
        [NYCPS],
    ),
    (
        "Can a teacher use a personal ChatGPT, Gemini, Copilot, or MagicSchool account?",
        "Only if that exact product and account type are approved for that use by the school. Never assume the consumer version is equivalent to the managed school version.",
        "Account ownership, tenant controls, data handling, retention, sharing, model training, connectors, and administrative visibility can differ across versions.",
        "If the approved environment is unknown, stay in paper/blueprint mode and ask leadership. Do not create a workaround account during the lab.",
        [NYCPS_ERMA],
    ),
    (
        "What if a teacher only wants to analyze their own class?",
        "A legitimate educational purpose may exist, but the tool and disclosure questions still remain.",
        "FERPA may allow school officials to access records needed for their professional responsibilities, subject to local policy. Sending those records to a vendor or AI feature is a separate issue that depends on the school's approved relationship and controls.",
        "Use school-approved reports and tools. Escalate any new AI processing workflow before entering live data.",
        [("FERPA school-official guidance", "https://studentprivacy.ed.gov/sites/default/files/resource_document/file/FERPA%20Guidance%20for%20School%20Officials%20on%20Student%20Health%20Records.pdf"), NYCPS_ERMA],
    ),
    (
        "What should happen if someone accidentally pastes student information into an unapproved tool?",
        "Stop. Do not continue transforming, summarizing, or reposting the information.",
        "The participant should follow the school's incident, deletion, and reporting procedures. Andrew should not investigate the record, repeat it in notes, or promise that deletion from the chat resolves the issue.",
        "Pause the exercise, ask the participant not to share the information with the room, and contact the named school privacy/data or technology owner immediately.",
        [NYSED, NYCPS_ERMA],
    ),
    (
        "Could these advanced questions become a follow-up PD?",
        "Yes. They are valuable, but they need a different promise, prerequisites, and leadership ownership.",
        "The beginner lab teaches student-neutral assistant design. A follow-up lab could teach approved learning-pattern review using fictional data first, then a school-owned workflow only after data, tool, purpose, group-size, access, review, retention, and escalation rules are established.",
        "Offer an advanced discovery and design session to leadership before offering teacher practice with any student-level or group-level data.",
        [NYCPS, NYSED],
    ),
]

for item in faq_items:
    add_faq(doc, *item)

# Part 7
add_page_break(doc)
doc.add_heading("Part 7 | Follow-up PD: Learning Pattern Review", level=1)
doc.add_paragraph("This is a separate advanced product, not an extension teachers casually unlock after the beginner lab. It begins with leadership and data-governance decisions, then uses fictional or synthetic data before any live workflow is considered.")

doc.add_heading("Proposed promise", level=2)
add_callout(doc, "Advanced PD concept", "Help educators inspect approved learning-pattern data without turning AI into a grader, profiler, placement engine, or substitute for professional judgment.")

doc.add_heading("Required prerequisites", level=2)
for text in (
    "The exact school-approved product, account, feature, and data source are named.",
    "The school privacy/data owner approves the fields, purpose, access, retention, and export path.",
    "Leadership defines permitted instructional decisions and prohibited high-impact uses.",
    "Minimum group-size or small-cell suppression rules are defined.",
    "The workflow begins with fictional or synthetic data and passes a documented test.",
    "Educators receive bias, re-identification, interpretation, and escalation practice.",
    "Human review and documentation are required before any instructional response.",
):
    add_bullet(doc, text)

doc.add_heading("Possible 90-minute advanced sequence", level=2)
advanced = [
    ("0-15", "Purpose before data", "Name one instructional question and the decision that remains human."),
    ("15-30", "Data minimization", "Choose only the approved fields needed for that question."),
    ("30-45", "Identity and group risk", "Test direct, indirect, linkability, small-group, and proxy-variable risks."),
    ("45-60", "Fictional analysis", "Use a synthetic dataset to inspect patterns and failure modes."),
    ("60-75", "Interpretation and bias", "Compare the AI summary with source data and alternate explanations."),
    ("75-90", "Decision and escalation", "Document what the tool may surface, what educators decide, and when the workflow stops."),
]
table = doc.add_table(rows=1, cols=3)
headers = ["TIME", "STAGE", "OUTPUT"]
for idx, value in enumerate(headers):
    set_cell_fill(table.cell(0, idx), DEEP)
    r = table.cell(0, idx).paragraphs[0].add_run(value)
    set_run_font(r, "Arial Narrow", 9, WHITE, True)
set_repeat_table_header(table.rows[0])
for time, stage, output in advanced:
    cells = table.add_row().cells
    cells[0].text = time
    cells[1].text = stage
    cells[2].text = output
set_table_geometry(table, [1200, 2650, 6030], indent_dxa=0)

doc.add_heading("What Andrew can responsibly own", level=2)
for text in (
    "Facilitation, task framing, workflow design, fictional practice, human review routines, and measurement.",
    "Making hidden assumptions, decision lines, and failure modes visible.",
    "Creating platform adapters after the school names its approved environment.",
):
    add_bullet(doc, text)

doc.add_heading("What remains with school specialists", level=2)
for text in (
    "Legal interpretation, compliance certification, privacy/security approval, procurement, and vendor contracts.",
    "Data-access authorization, de-identification standards, key management, retention, deletion, and incident response.",
    "Special-education, counseling, discipline, safety, placement, and other regulated decisions.",
):
    add_bullet(doc, text)

doc.add_heading("Leadership discovery before offering this", level=2)
for text in (
    "What instructional question are educators trying to answer, and what decision must remain human?",
    "Which exact product, account, feature, data source, and export path are already approved?",
    "Who owns privacy/data, technology, instructional policy, and regulated specialist decisions?",
    "Which data fields are necessary, which are prohibited, and what group-size or suppression rule applies?",
    "What output may be retained, who may access it, when is it deleted, and how are uncertain findings escalated?",
    "What evidence would make leadership continue, revise, or stop the pilot?",
):
    add_bullet(doc, text)

add_callout(doc, "Proceed only if", "Leadership can answer every discovery question, name the responsible owners, and authorize a fictional-data rehearsal before any live-data workflow is considered.")

# Appendix quick ladder
add_page_break(doc)
doc.add_heading("Appendix A | Six-gate student-data decision ladder", level=1)
doc.add_paragraph("Use this ladder when a participant asks whether AI may process student work, records, coded rows, groups, or learning trends. One unknown answer means stop and escalate. It is not a compliance checklist.")

gates = [
    ("1", "PURPOSE", "Is the instructional purpose legitimate, specific, and necessary?"),
    ("2", "ENVIRONMENT", "Is the exact product, account, and feature approved for this use?"),
    ("3", "DATA", "Are the fields and records approved, minimum necessary, and handled through the authorized route?"),
    ("4", "IDENTITY", "Has the school, not the teacher alone, addressed direct, indirect, coded, linked, and small-group identification risk?"),
    ("5", "DECISION", "Does the educator retain the decision, with high-impact uses prohibited or separately governed?"),
    ("6", "REVIEW", "Are accuracy, bias, interpretation, documentation, retention, and escalation defined?"),
]
for number, label, question in gates:
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [700, 1700, 7480], indent_dxa=0)
    set_cell_fill(table.cell(0, 0), GOLD)
    set_cell_fill(table.cell(0, 1), DEEP)
    set_cell_fill(table.cell(0, 2), "F8FAF9")
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(number)
    set_run_font(r0, "Arial Narrow", 15, DEEP, True)
    r1 = table.cell(0, 1).paragraphs[0].add_run(label)
    set_run_font(r1, "Arial Narrow", 9.5, WHITE, True)
    r2 = table.cell(0, 2).paragraphs[0].add_run(question)
    set_run_font(r2, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

add_callout(doc, "If any gate is unknown", "Do not enter live data. Continue with fictional material or paper planning. Record the question and route it to the named school owner.", fill=RISK, accent="A13D2D")

doc.add_heading("Appendix B | Calm facilitator language", level=1)
doc.add_heading("Suggested responses", level=2)
responses = [
    ("I do not understand AI.", "That is fine. You need one repeated task and your professional judgment. The structure does the rest."),
    ("Can I remove the name and use it?", "A name is only one identifier. For this lab, use fictional or student-neutral material. We can record the real workflow as a question for the school's data/privacy owner."),
    ("But this is my own class.", "Your instructional reason may be legitimate. The separate question is whether this AI product, account, feature, and data path are approved. We will not self-authorize that in this workshop."),
    ("The vendor says it is compliant.", "That may be useful information for the school review, but it does not replace the school's approval of the exact tool and use."),
    ("Can it make groups for me?", "It can help us examine a fictional workflow. Live grouping affects students and requires a separately governed process with educator ownership."),
    ("What if the approved tool is unavailable?", "We stay in paper/demo mode. You can finish the blueprint without creating a workaround account."),
    ("What if it takes longer than doing it myself?", "Then revise once or stop. The point is recovered time, not AI use for its own sake."),
]
for prompt, response in responses:
    doc.add_heading(prompt, level=3)
    p = doc.add_paragraph(response)
    p.paragraph_format.left_indent = Inches(0.18)

doc.add_heading("Appendix C | Official sources", level=1)
doc.add_paragraph("Last verified August 30, 2026. Local policy, contracts, platform capabilities, and official guidance can change. Recheck before each advanced-data engagement.")
sources = [
    ("U.S. Department of Education: Personally Identifiable Information for Education Records", FERPA_PII[1]),
    ("U.S. Department of Education: What Constitutes De-identified Records and Information?", FERPA_DEID[1]),
    ("U.S. Department of Education: FERPA Regulations", FERPA_RULE[1]),
    ("New York State Education Department: Education Law Section 2-d Resources", NYSED[1]),
    ("New York City Public Schools: Guidance on Artificial Intelligence", NYCPS[1]),
    ("New York City Public Schools: AI Tool Evaluation and ERMA", NYCPS_ERMA[1]),
]
for label, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, label, url)

doc.add_paragraph()
add_callout(doc, "Final reminder", "Andrew teaches the build method. School leaders name the approved environment. Qualified specialists resolve regulated questions. Teachers retain professional judgment and every final decision.", fill=DEEP, accent=AQUA)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)

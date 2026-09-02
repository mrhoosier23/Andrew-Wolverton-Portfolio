from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import BaseDocTemplate, Frame, HRFlowable, Image, PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
PDF_DIR = ROOT / "output" / "pdf"
DOCX_DIR = ROOT / "output" / "docx"
ASSET_DIR = ROOT / "flagship-demo" / "assets"
PDF_DIR.mkdir(parents=True, exist_ok=True)
DOCX_DIR.mkdir(parents=True, exist_ok=True)

DEEP = "#063C33"
TEAL = "#008F87"
IVORY = "#F6F0E5"
CHARCOAL = "#1D2A27"
GOLD = "#E4A11B"
MINT = "#E4F5F1"
LINE = "#C9D5D0"
WHITE = "#FFFDF8"
MUTED = "#4D6761"
AUBERGINE = "#50323F"


def register_fonts():
    regular = Path(r"C:\Windows\Fonts\arial.ttf")
    bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    narrow = Path(r"C:\Windows\Fonts\arialnb.ttf")
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("AWBody", str(regular)))
        pdfmetrics.registerFont(TTFont("AWBold", str(bold)))
        if narrow.exists():
            pdfmetrics.registerFont(TTFont("AWDisplay", str(narrow)))
            return "AWBody", "AWBold", "AWDisplay"
        return "AWBody", "AWBold", "AWBold"
    return "Helvetica", "Helvetica-Bold", "Helvetica-Bold"


BODY_FONT, BOLD_FONT, DISPLAY_FONT = register_fonts()
sample = getSampleStyleSheet()
PDF_TITLE = ParagraphStyle("PDFTitle", parent=sample["Title"], fontName=DISPLAY_FONT, fontSize=29, leading=30, textColor=colors.HexColor(DEEP), alignment=TA_LEFT, spaceAfter=10)
PDF_H1 = ParagraphStyle("PDFH1", parent=sample["Heading1"], fontName=DISPLAY_FONT, fontSize=20, leading=22, textColor=colors.HexColor(DEEP), spaceBefore=9, spaceAfter=7)
PDF_H2 = ParagraphStyle("PDFH2", parent=sample["Heading2"], fontName=BOLD_FONT, fontSize=13.5, leading=16, textColor=colors.HexColor(TEAL), spaceBefore=7, spaceAfter=5)
PDF_BODY = ParagraphStyle("PDFBody", parent=sample["BodyText"], fontName=BODY_FONT, fontSize=11, leading=15, textColor=colors.HexColor(CHARCOAL), spaceAfter=7)
PDF_SMALL = ParagraphStyle("PDFSmall", parent=PDF_BODY, fontSize=9.5, leading=13, textColor=colors.HexColor(MUTED))
PDF_LABEL = ParagraphStyle("PDFLabel", parent=PDF_BODY, fontName=BOLD_FONT, fontSize=10, leading=12, textColor=colors.HexColor(TEAL), spaceAfter=5)
PDF_BIG = ParagraphStyle("PDFBig", parent=PDF_TITLE, fontSize=24, leading=25)


def P(text: str, style=PDF_BODY):
    return Paragraph(text, style)


def pdf_bullet(text: str):
    return P(f"<font color='{TEAL}'><b>●</b></font>&nbsp;&nbsp;{text}")


def pdf_checkbox(text: str):
    return P(f"<font color='{TEAL}'><b>□</b></font>&nbsp;&nbsp;{text}")


def pdf_rule():
    return HRFlowable(width="100%", thickness=0.8, color=colors.HexColor(LINE), spaceBefore=6, spaceAfter=9)


def note_box(label: str, body: str, fill=MINT, accent=TEAL):
    table = Table([[P(label.upper(), PDF_LABEL), P(body)]], colWidths=[1.45 * inch, 4.75 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(fill)),
        ("LINEBEFORE", (0, 0), (0, 0), 6, colors.HexColor(accent)),
        ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 11),
        ("RIGHTPADDING", (0, 0), (-1, -1), 11),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    return table


def writing_box(label: str, height=0.65 * inch):
    table = Table([[P(label, PDF_LABEL)], [""]], colWidths=[6.15 * inch], rowHeights=[0.24 * inch, height])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(WHITE)),
        ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    return table


def pdf_footer(canvas, doc):
    canvas.saveState()
    w, h = doc.pagesize
    canvas.setFillColor(colors.HexColor(IVORY))
    canvas.rect(0, 0, w, h, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor(DEEP))
    canvas.rect(0, h - 0.12 * inch, w, 0.12 * inch, fill=1, stroke=0)
    canvas.setFont(BOLD_FONT, 8)
    canvas.setFillColor(colors.HexColor(TEAL))
    canvas.drawString(doc.leftMargin, 0.34 * inch, "TEACHER TIME BACK LAB")
    canvas.setFont(BODY_FONT, 8)
    canvas.setFillColor(colors.HexColor(MUTED))
    canvas.drawRightString(w - doc.rightMargin, 0.34 * inch, f"Andrew Wolverton  |  {doc.page}")
    canvas.restoreState()


def build_pdf(path: Path, story, pagesize=letter, margins=(0.72, 0.72, 0.72, 0.65)):
    left, right, top, bottom = [x * inch for x in margins]
    doc = BaseDocTemplate(str(path), pagesize=pagesize, leftMargin=left, rightMargin=right, topMargin=top, bottomMargin=bottom, title=path.stem.replace("-", " "), author="Andrew Wolverton")
    frame = Frame(left, bottom, pagesize[0] - left - right, pagesize[1] - top - bottom, id="body")
    doc.addPageTemplates([PageTemplate(id="standard", frames=[frame], onPage=pdf_footer)])
    doc.build(story)


def pdf_cover(kicker: str, title: str, deck: str):
    return [P(kicker.upper(), PDF_LABEL), P(title, PDF_TITLE), P(deck), Spacer(1, 0.08 * inch)]


def image_flow(name: str, width: float, height: float):
    return Image(str(ASSET_DIR / name), width=width, height=height)


FIVE_FIELDS = [
    ("1", "Help me with", "Name one repeated job, not a broad role."),
    ("2", "Use only", "Name the exact blank format, rough notes, public standard, or other permitted material for this task."),
    ("3", "Return", "Describe the useful shape of the first draft."),
    ("4", "Stop and ask if", "Name what must be present. Missing information becomes NEEDS TEACHER INPUT."),
    ("5", "I will review", "Check facts, bias, appropriateness, alignment, missing information, and final wording."),
]

LEADERSHIP_AGENDA = [
    ("0–5", "Frame the decision", "Name one exact product, one staff account type, and one school contact."),
    ("5–15", "Name the tool and account", "Confirm login path, account controls, and cohort access."),
    ("15–25", "Set the boundary", "No real student information or work. Confirm local rules and escalation."),
    ("25–35", "Confirm operations", "Cohort, space, devices, accessibility, technology support, and non-evaluative participation."),
    ("35–42", "Preview the product", "Show the Weekly Lesson Setup Assistant cycle."),
    ("42–45", "Lock next actions", "Complete the Platform Setup Record."),
]

LAB_AGENDA = [
    ("0–10", "See the finished product"),
    ("10–22", "Choose one repeated task"),
    ("22–32", "Check the school tool and account"),
    ("32–52", "Build the five fields"),
    ("52–68", "Practice and review"),
    ("68–78", "Revise once and retest"),
    ("78–86", "Save and begin the two-week log"),
    ("86–90", "Exit check and next step"),
]

FOLLOWUP_AGENDA = [
    ("0–5", "Restate the decision", "Decide whether each workflow should be kept, revised, or stopped."),
    ("5–15", "Review the evidence", "Uses, before time, assisted time, correction and review time, and result quality."),
    ("15–23", "Troubleshoot patterns", "Discuss broad tasks, missing inputs, weak output shape, and review burden without exposing participant content."),
    ("23–28", "Make decisions", "Each teacher chooses keep, revise, or stop."),
    ("28–30", "Close", "Leadership receives only cohort-level learning."),
]

QUIZ = [
    ("Which account should a teacher use?", "The exact AI tool and staff account the school named.", "A personal or workaround account is not the lab path."),
    ("Is replacing names with Student A or a key ID automatically safe?", "No.", "Other details, group size, context, and linked records may still identify a student."),
    ("What should happen when required information is missing?", "The assistant marks NEEDS TEACHER INPUT.", "It does not guess or invent the missing detail."),
    ("Who makes the final decision?", "The teacher.", "The assistant prepares. The teacher reviews, corrects, and decides."),
]

FAQS = [
    ("Can a teacher review specific student work if names are removed?", "Not in the introductory lab. Names are only one kind of identifier. Initials, key IDs, small groups, dates, needs, quotations, and surrounding context may still identify a student. A later engagement requires school review of the platform, purpose, fields, access, retention, reporting, and specialist responsibilities before student work is used."),
    ("Can teachers track groups or learning trends?", "Not in this introductory lab. A later workflow may be possible only after the school defines the educational purpose, minimum necessary data, access, retention, reporting, and decision ownership. Small groups can create re-identification risk. Andrew can facilitate workflow design, but school privacy, legal, security, special-education, and instructional leaders make the applicable decisions."),
    ("What if a separate key maps coded IDs back to students?", "A separate key can reduce casual exposure, but it does not automatically de-identify data. If someone can reasonably reconnect the code to a student, or surrounding details identify the student, treat it as identifiable. This requires a separate school-reviewed data workflow."),
    ("Can the assistant grade, group, place, diagnose, discipline, or recommend accommodations?", "No. Those decisions remain with teachers and qualified school personnel. The introductory lab focuses on bounded drafting and organizing work that a teacher reviews."),
    ("Can teachers use publisher curriculum or a colleague's materials?", "Only when the school and platform permit it. The lab does not assume licensed or third-party materials may be uploaded."),
    ("What if a teacher has no AI account?", "They do not create a workaround or personal account. The school resolves access. A teacher may observe and draft on paper, but has not completed the product until the setup is saved and tested in the permitted account."),
    ("Is this only for OpenAI products?", "No. The five-part method can be adapted to ChatGPT, Microsoft Copilot, Google Gemini, MagicSchool, or another school-selected product. Interfaces and controls differ, so Andrew prepares and tests one platform setup card for the actual account type before the pilot."),
    ("Does the lab guarantee time savings?", "No. Teachers measure total time, including correction and review, for two weeks and choose keep, revise, or stop."),
    ("Is this compliance or legal training?", "No. Andrew teaches a bounded workflow and safety routine. The school and qualified specialists retain policy, legal, privacy, security, procurement, labor, special-education, and other regulated decisions."),
    ("What are Andrew's credentials?", "Andrew Wolverton is a former Indiana classroom educator, program builder, facilitator, and human-led systems designer. He helps teams turn confusing workflows into practical tools people can understand, review, and use."),
]


def pilot_overview_pdf():
    story = pdf_cover("Founding-school pilot", "Give teachers back time. Build one useful AI assistant in 90 minutes.", "A practical, beginner-friendly lab for schools that want AI to reduce work, not add another technology burden.")
    left = [P("Bring one recurring task. Leave with one assistant you can use again.", PDF_H2), P("Andrew demonstrates a finished workflow, then guides each teacher through building and testing reusable instructions inside the AI tool and staff account the school already allows."), pdf_bullet("No coding and no previous AI experience"), pdf_bullet("No real student information, student work, or confidential records"), pdf_bullet("One completed setup, one practice result, and one revision"), pdf_bullet("A two-week keep, revise, or stop decision")]
    hero = Table([[left, image_flow("05-before-after.png", 2.82 * inch, 2.10 * inch)]], colWidths=[3.35 * inch, 2.85 * inch])
    hero.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 8)]))
    story.extend([hero, Spacer(1, 0.09 * inch), pdf_rule()])
    facts = [[P("45 + 90 + 30", PDF_BIG), P("Leadership setup, teacher lab, and follow-up", PDF_SMALL)], [P("8–20", PDF_BIG), P("Educators in one cohort", PDF_SMALL)], [P("$1,500", PDF_BIG), P("Founding pilot; second same-day cohort $750", PDF_SMALL)]]
    t = Table([facts], colWidths=[2.07 * inch] * 3)
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(MINT)), ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor(LINE)), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))
    story.extend([t, Spacer(1, 0.12 * inch), note_box("School readiness", "Before booking, leadership names the exact AI product and staff account, the local rules, and one school contact for questions.", fill="#FFF4D7", accent=GOLD), Spacer(1, 0.08 * inch), P("Schedule a short pilot-fit conversation. Do not send student information or confidential school records in the inquiry.", PDF_H2)])
    build_pdf(PDF_DIR / "Teacher-Time-Back-Lab-Pilot-Overview.pdf", story)


def leadership_pdf():
    story = pdf_cover("For school leadership", "Leadership readiness guide", "Use this during the 45-minute setup conversation so the teacher lab is concrete, safe, and finishable before anyone enters the room.")
    story.extend([note_box("The decision", "Name one exact AI product and staff account for the lab. If the school has not selected one, pause and hold a separate readiness conversation.", fill="#FFF4D7", accent=GOLD), P("What leadership confirms", PDF_H1)])
    for item in ["Exact AI product, staff account type, and login path", "Current staff-use and data-handling guidance", "Whether file upload, sharing, browsing, connectors, and memory are allowed", "A school contact for policy or platform questions", "8–20 participants with working accounts and laptops", "Internet, display, accessibility, and technology support", "Results will not evaluate individual teachers"]:
        story.append(pdf_checkbox(item))
    story.extend([P("The introductory boundary", PDF_H1), P("Teachers use practice copy with no real people or records. Replacing names with Student A, initials, group labels, or key IDs does not automatically make information safe."), PageBreak(), P("Who decides what", PDF_H1)])
    rows = [[P("Andrew", PDF_H2), P("Demonstrates the workflow, teaches the method, leads practice, and keeps human review visible.")], [P("School leadership", PDF_H2), P("Selects the tool and account, owns local policy and permissions, and identifies the escalation contact.")], [P("Teachers", PDF_H2), P("Choose a bounded task, protect information, review every result, and make final decisions.")], [P("Qualified specialists", PDF_H2), P("Handle privacy, legal, security, special-education, labor, procurement, and regulated questions.")]]
    tbl = Table(rows, colWidths=[1.55 * inch, 4.65 * inch])
    tbl.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(MINT))]))
    story.extend([tbl, P("Technology fallback", PDF_H1), P("Andrew can demonstrate offline and teachers can complete the paper setup if access fails. The product is not complete until the setup is saved and tested in the permitted account."), P("Leadership follow-up", PDF_H1), P("Leadership receives cohort-level counts and medians, not participant names, task content, or identifiable comments."), writing_box("School tool and account"), writing_box("School contact for questions"), writing_box("Local rules or links")])
    build_pdf(PDF_DIR / "Leadership-Readiness-Guide.pdf", story)


def prep_pdf():
    story = pdf_cover("For participants", "What to bring to the Teacher Time Back Lab", "You do not need to know AI. Bring one repeated task and a working login. Andrew shows the finished result before you build anything.")
    story.append(P("Bring", PDF_H1))
    for item in ["A laptop and charger", "A working login to the exact AI tool and staff account your school named", "One recurring task you personally perform", "Optional rough notes or a blank format you created or are permitted to reuse"]:
        story.append(pdf_checkbox(item))
    story.append(P("Do not bring or upload", PDF_H1))
    for item in ["Student work, grades, rosters, IDs, initials, group labels, or case notes", "Family messages, staff records, personnel matters, or confidential school documents", "Publisher curriculum or colleague-created resources unless permitted", "Passwords or account access for Andrew"]:
        story.append(pdf_bullet(item))
    story.extend([note_box("No task yet?", "Choose lesson setup from rough notes, directions formatter, or meeting agenda and follow-up organizer."), note_box("What you leave with", "One saved five-part setup, one practice result, one revision, a review card, and a two-week plan.", fill="#FFF4D7", accent=GOLD)])
    build_pdf(PDF_DIR / "Teacher-Preparation-Sheet.pdf", story)


def workbook_pdf():
    story = pdf_cover("Participant workbook", "Build one assistant you can use again", "Follow six workshop steps. Write plainly. The assistant prepares. The teacher decides.")
    story.extend([note_box("Lab boundary", "Use no real student information, student work, family records, personnel records, or confidential school records."), P("The six steps", PDF_H1)])
    for n, title in enumerate(["See the finished product", "Choose one repeated task", "Check the school tool and account", "Build the five fields", "Practice and review", "Save and measure"], 1):
        story.append(P(f"<font color='{GOLD}'><b>{n:02d}</b></font>&nbsp;&nbsp;<b>{title}</b>"))
    story.extend([PageBreak(), P("Step 1  |  See the finished product", PDF_H1), image_flow("05-before-after.png", 6.1 * inch, 4.55 * inch), P("Rough notes → saved instructions → first result → teacher correction → reuse."), PageBreak(), P("Step 2  |  Choose one repeated task", PDF_H1), P("A strong first task repeats, takes meaningful time, produces a reviewable draft, and can be practiced without real student or confidential information."), writing_box("My repeated task", 1.0 * inch), writing_box("How often I do it", 0.55 * inch), writing_box("Estimated minutes now", 0.55 * inch), note_box("Starter choices", "Lesson setup from rough notes • Directions formatter • Meeting agenda and follow-up organizer"), PageBreak(), P("Step 3  |  Check the school tool and account", PDF_H1), writing_box("Exact AI product", 0.55 * inch), writing_box("Staff account or login path", 0.55 * inch), writing_box("School contact if I am unsure", 0.55 * inch), *[pdf_checkbox(x) for x in ["I am in the school-named tool and account.", "My practice copy contains no real people or records.", "I know what I will review before anything is used."]], PageBreak(), P("Step 4  |  Build the five fields", PDF_H1)])
    for n, label, help_text in FIVE_FIELDS:
        story.extend([P(f"<font color='{GOLD}'><b>{n}</b></font>&nbsp;&nbsp;<b>{label}</b>  <font color='{MUTED}'>{help_text}</font>"), writing_box(label, 0.62 * inch)])
    story.extend([PageBreak(), P("Step 5  |  Practice and review", PDF_H1), note_box("Exact source boundary", "Use only the text or files I give you for this task. If the information is not there, mark NEEDS TEACHER INPUT instead of guessing.", fill="#FFF4D7", accent=GOLD), *[pdf_checkbox(x) for x in ["Facts and calculations are correct.", "The result is appropriate for the audience and purpose.", "Bias, assumptions, or unfair framing have been addressed.", "Missing information is visible instead of invented.", "Final wording and instructional choices remain mine."]], writing_box("One correction I made", 0.75 * inch), PageBreak(), P("Step 6  |  Save and measure", PDF_H1), writing_box("Saved assistant name", 0.55 * inch), writing_box("Where I saved it", 0.55 * inch)])
    data = [[P("Date", PDF_LABEL), P("Task", PDF_LABEL), P("Before", PDF_LABEL), P("Using", PDF_LABEL), P("Review", PDF_LABEL), P("Result", PDF_LABEL)]] + [["", "", "", "", "", ""] for _ in range(5)]
    log = Table(data, colWidths=[0.75 * inch, 1.65 * inch, 0.78 * inch, 0.78 * inch, 0.78 * inch, 1.0 * inch], rowHeights=[0.32 * inch] + [0.48 * inch] * 5)
    log.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(MINT)), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.extend([log, Spacer(1, 0.10 * inch), note_box("Decision", "KEEP if useful and total time drops. REVISE if promising but the setup needs work. STOP if risk or review burden outweighs the benefit.")])
    build_pdf(PDF_DIR / "Teacher-Workbook.pdf", story)


def practice_pack_pdf():
    story = pdf_cover("Practice pack", "Weekly Lesson Setup Assistant", "A complete example with no real people or records. Practice before using your own permitted materials.")
    items = [("1. Rough notes", "01-rough-notes.png", "The teacher begins with topic, timing, materials, and activity ideas."), ("2. Saved instructions", "02-saved-instructions.png", "The five fields define the job, inputs, output, stop rule, and teacher review."), ("3. First result", "03-first-result.png", "The draft follows the weekly format and marks missing information."), ("4. Teacher revision", "04-teacher-revision.png", "The teacher corrects one issue and decides whether the result is usable.")]
    for i, (title, asset, body) in enumerate(items):
        story.extend([P(title, PDF_H1), image_flow(asset, 6.05 * inch, 4.54 * inch), P(body)])
        if i != len(items) - 1:
            story.append(PageBreak())
    build_pdf(PDF_DIR / "Weekly-Lesson-Setup-Practice-Pack.pdf", story)


def safety_card_pdf():
    page = landscape(letter)
    story = pdf_cover("Desk reference", "Stop, ask, or continue", "Use the label and symbol, not color alone. The introductory lab never uses real student information or student work.")
    cards = [("[READY]  READY FOR THE LAB", "Practice copy with no real people or records; a permitted blank format or rough notes; a public standard when permitted.", MINT, TEAL), ("[ASK]  ASK SCHOOL BEFORE LATER USE", "Real student work, coded or grouped data, licensed curriculum, family communication, connectors, file uploads, sharing, or memory.", "#FFF4D7", GOLD), ("[STOP]  NOT USED IN THIS LAB", "Grades, rosters, IEPs, behavior notes, placement, discipline, accommodations, evaluations, or student-level decisions.", "#EEE8EA", AUBERGINE)]
    row = []
    for title, body, fill, accent in cards:
        box = Table([[P(title, PDF_H2)], [P(body)]], colWidths=[2.9 * inch], rowHeights=[0.48 * inch, 1.28 * inch])
        box.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(fill)), ("BOX", (0, 0), (-1, -1), 1.1, colors.HexColor(accent)), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 12), ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))
        row.append(box)
    table = Table([row], colWidths=[3.05 * inch] * 3)
    table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.extend([table, Spacer(1, 0.16 * inch), P("The five-question review", PDF_H1)])
    review = [[P(f"<font color='{GOLD}'><b>{i}</b></font>", PDF_BIG), P(text)] for i, text in enumerate(["Is this the school-named tool and staff account?", "Does the material contain information about a real student, family, staff member, or confidential situation?", "Did the assistant stay inside the material and mark missing information?", "Are the facts, tone, fairness, and instructional choices appropriate?", "Am I prepared to own the final decision and wording?"], 1)]
    rt = Table(review, colWidths=[0.55 * inch, 8.5 * inch])
    rt.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -2), 0.6, colors.HexColor(LINE)), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
    story.extend([rt, note_box("When unsure", "Stop. Do not enter the material. Ask the school contact.", fill="#FFF4D7", accent=GOLD)])
    build_pdf(PDF_DIR / "Safety-and-Review-Card.pdf", story, pagesize=page, margins=(0.55, 0.55, 0.55, 0.55))


def platform_record_pdf():
    story = pdf_cover("Pilot setup record", "The exact tool and account for this cohort", "Complete with the school host before the lab. This is a local decision record, not a universal platform approval.")
    for label in ["School or program", "Pilot date", "Exact AI product", "Staff account type and login path", "School host and escalation contact"]:
        story.append(writing_box(label, 0.46 * inch))
    story.extend([PageBreak(), P("Feature decisions", PDF_H1)])
    rows = [[P("Feature", PDF_LABEL), P("Allowed?", PDF_LABEL), P("School note", PDF_LABEL)]] + [[P(x), "□ Yes   □ No   □ Ask", ""] for x in ["Pasting practice text", "Uploading files", "Web browsing or search", "Memory or chat history", "Sharing assistants", "Connectors or external apps", "Saving participant work"]]
    tbl = Table(rows, colWidths=[2.0 * inch, 1.55 * inch, 2.65 * inch], rowHeights=[0.42 * inch] + [0.55 * inch] * 7)
    tbl.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(MINT)), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.extend([tbl, note_box("Minimum first-lab setup", "The first lab can run without uploads, browsing, connectors, or sharing. Teachers can paste practice text and save instructions."), P("Test before the session", PDF_H1), *[pdf_checkbox(x) for x in ["Participant login works on the school network.", "The five-part setup can be saved or copied.", "Practice text can be entered without unapproved features.", "The facilitator can demonstrate in the same account type.", "The school host knows how participants request help."]]])
    build_pdf(PDF_DIR / "Platform-Setup-Record.pdf", story)


def facilitator_pdf():
    story = pdf_cover("Andrew and approved co-facilitators", "Teacher Time Back Lab facilitator playbook", "Leadership setup, 90-minute teacher lab, technology fallback, nuanced FAQ, and follow-up.")
    story.extend([note_box("Core promise", "Bring one task you keep doing. Leave with one assistant you can use again."), note_box("Human line", "The assistant prepares. The teacher decides.", fill="#FFF4D7", accent=GOLD), P("Before outreach", PDF_H1)])
    for item in ["Sales page and flagship demo are current", "Pilot overview is attached, not the full curriculum", "Leadership setup happens before the lab", "No guaranteed time-saving, compliance, or licensure claim", "Andrew receives no passwords or student data"]:
        story.append(pdf_checkbox(item))
    story.extend([PageBreak(), P("Leadership setup  |  45 minutes", PDF_H1)])
    lt = Table([[P("Time", PDF_LABEL), P("Purpose", PDF_LABEL), P("Say and do", PDF_LABEL)]] + [[P(a), P(b, PDF_H2), P(c)] for a, b, c in LEADERSHIP_AGENDA], colWidths=[0.65 * inch, 1.55 * inch, 4.0 * inch])
    lt.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(MINT)), ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.extend([lt, PageBreak(), P("Opening script  |  first 10 minutes", PDF_H1), P("<b>Welcome:</b> Today is not a tour of AI. You are going to stop rebuilding one piece of work. I will show you a finished example first, then we will build your version together. You do not need coding or previous AI experience."), P("<b>Boundary:</b> We will not use real student information, student work, family records, personnel records, or confidential school files. If you are unsure, stop and ask the school contact."), P("<b>Proof:</b> Show rough notes, saved instructions, first result, and teacher correction before explaining the fields."), P("<b>Responsibility:</b> The assistant prepares. The teacher decides."), PageBreak(), P("Teacher lab  |  90 minutes", PDF_H1)])
    at = Table([[P("Time", PDF_LABEL), P("Participant action", PDF_LABEL), P("Proof before moving on", PDF_LABEL)]] + [[P(a), P(b, PDF_H2), P("Visible completion")] for a, b in LAB_AGENDA], colWidths=[0.75 * inch, 2.55 * inch, 2.9 * inch])
    at.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(MINT)), ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.extend([at, PageBreak(), P("Step prompts and checks", PDF_H1)])
    scripts = [("Choose one repeated task", "What do you rebuild often enough to resent, but can still review before using?", "One bounded drafting or organizing job."), ("Check the tool and account", "Read the product and account from the school record. If yours does not match, stop and ask.", "Correct account."), ("Build the five fields", "Write one sentence for each field. Specific beats clever.", "Five fields and source boundary."), ("Practice", "Use the provided practice pack first.", "One practice result."), ("Review", "Find one thing you would change.", "One teacher correction."), ("Save and measure", "Name it for next week. Count correction and review time.", "Saved setup and understood log.")]
    for title, say, proof in scripts:
        story.extend([P(title, PDF_H2), P(f"<b>Say:</b> {say}"), P(f"<b>Check:</b> {proof}")])
    story.extend([PageBreak(), P("Safety rehearsal  |  four questions", PDF_H1), P("Say: This is a four-question practice. Choose one answer. When the safe answer is selected, the visual changes and explains why.")])
    for i, (q, a, why) in enumerate(QUIZ, 1):
        story.extend([P(f"{i}. {q}", PDF_H2), P(f"<b>Correct:</b> {a} {why}")])
    story.extend([PageBreak(), P("Technology fallback", PDF_H1), P("Continue from exported images and complete the five fields on paper. The school host handles access. Do not create workaround accounts. Schedule completion because the product is not finished until saved and tested."), P("Two-week follow-up  |  30 minutes", PDF_H1)])
    ft = Table([[P("Time", PDF_LABEL), P("Move", PDF_LABEL), P("Language", PDF_LABEL)]] + [[P(a), P(b, PDF_H2), P(c)] for a, b, c in FOLLOWUP_AGENDA], colWidths=[0.65 * inch, 1.45 * inch, 4.1 * inch])
    ft.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(MINT)), ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(LINE)), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(LINE)), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.extend([ft, PageBreak(), P("Nuanced FAQ  |  student information and instructional boundaries", PDF_H1)])
    for index, (q, a) in enumerate(FAQS):
        if index == 5:
            story.extend([PageBreak(), P("Nuanced FAQ  |  accounts, platforms, claims, and credentials", PDF_H1)])
        story.extend([P(q, PDF_H2), P(a)])
    story.extend([PageBreak(), P("Completion check", PDF_H1), *[pdf_checkbox(x) for x in ["Task is bounded and repeated", "Tool and account match the school record", "Five fields and source boundary are complete", "Practice used no real people or records", "One result was reviewed and revised", "Setup is saved", "Teacher understands the two-week decision"]], P("Attendance alone is not completion.", PDF_H2)])
    build_pdf(PDF_DIR / "Teacher-Time-Back-Lab-Facilitator-Playbook.pdf", story)


def set_docx_font(run, size=11, bold=False, color="000000", name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_google_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after, color in [("Heading 1", 20, 20, 6, "000000"), ("Heading 2", 16, 18, 6, "000000"), ("Heading 3", 14, 16, 4, "434343")]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = False
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for list_name in ["List Bullet", "List Number"]:
        st = doc.styles[list_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15


def add_doc_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    set_docx_font(p.add_run(title), size=26)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_docx_font(p.add_run(subtitle), size=11, color="555555")


def add_doc_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_properties.append(header_marker)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_docx_font(run, size=10.5, bold=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F1F3F4")
        cell._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for run in cells[i].paragraphs[0].runs:
                set_docx_font(run, size=10.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def save_google_doc(path: Path, title: str, subtitle: str, sections):
    doc = Document()
    style_google_doc(doc)
    add_doc_title(doc, title, subtitle)
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        for kind, value in content:
            if kind == "p":
                doc.add_paragraph(value)
            elif kind == "h2":
                doc.add_heading(value, level=2)
            elif kind == "bullets":
                for item in value:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "numbered":
                for item in value:
                    doc.add_paragraph(item, style="List Number")
            elif kind == "table":
                headers, rows, widths = value
                add_doc_table(doc, headers, rows, widths)
    doc.core_properties.title = title
    doc.core_properties.author = "Andrew Wolverton"
    doc.core_properties.subject = "Teacher Time Back Lab"
    doc.save(path)


def build_docx_suite():
    facilitator_sections = [
        ("Purpose and promise", [("p", "Bring one task you keep doing. Leave with one assistant you can use again."), ("p", "The assistant prepares. The teacher decides."), ("bullets", ["45-minute leadership setup", "90-minute teacher lab", "30-minute follow-up two weeks later", "8–20 educators", "$1,500 founding-school pilot; second same-day cohort $750"])]),
        ("Leadership setup", [("table", (["Time", "Purpose", "Action"], LEADERSHIP_AGENDA, [0.75, 1.65, 4.1]))]),
        ("Opening script", [("p", "Today is not a tour of AI. You are going to stop rebuilding one piece of work. I will show you a finished example first, then we will build your version together. You do not need coding or previous AI experience."), ("p", "We will not use real student information, student work, family records, personnel records, or confidential school files. If you are unsure, stop and ask the school contact."), ("p", "Show rough notes, saved instructions, first result, and teacher correction before explaining the five fields.")]),
        ("Teacher lab run of show", [("table", (["Time", "Participant action"], LAB_AGENDA, [1.0, 5.5]))]),
        ("The five-part setup", [("numbered", [f"{label}: {help_text}" for _, label, help_text in FIVE_FIELDS]), ("p", "Use only the text or files I give you for this task. If the information is not there, mark NEEDS TEACHER INPUT instead of guessing.")]),
        ("Four-question safety rehearsal", [("p", "This is a four-question practice. Choose one answer. When the safe answer is selected, the visual changes and explains why."), ("numbered", [f"{q} Correct answer: {a} {why}" for q, a, why in QUIZ])]),
        ("Technology fallback", [("p", "Continue from exported images and complete the five fields on paper. The school host handles access. Do not create workaround accounts. Schedule completion because the product is not finished until saved and tested.")]),
        ("Two-week follow-up", [("table", (["Time", "Move", "Language"], FOLLOWUP_AGENDA, [0.75, 1.45, 4.3]))]),
        ("Nuanced questions and answers", sum(([('h2', q), ('p', a)] for q, a in FAQS), [])),
        ("Completion check", [("bullets", ["Task is bounded and repeated", "Tool and account match the school record", "Five fields and source boundary are complete", "Practice used no real people or records", "One result was reviewed and revised", "Setup is saved", "Teacher understands the two-week decision"]), ("p", "Attendance alone is not completion.")]),
    ]
    save_google_doc(DOCX_DIR / "Teacher-Time-Back-Lab-Facilitator-Playbook.docx", "Teacher Time Back Lab Facilitator Playbook", "Delivery script, run of show, technology fallback, and nuanced FAQ", facilitator_sections)

    leadership_sections = [
        ("The readiness decision", [("p", "Before booking the teacher lab, name one exact AI product and staff account. If the school has not selected one, hold a separate readiness conversation first.")]),
        ("Confirm before the lab", [("bullets", ["Exact AI product, staff account type, and login path", "Current local guidance for staff AI use and data handling", "A school contact for policy or platform questions", "8–20 participants with working accounts and laptops", "Internet, display, accessibility, and technology support", "Results will not evaluate individual teachers"])]),
        ("Introductory boundary", [("p", "No real student information or student work is used. Codes, initials, group labels, and separate key IDs do not automatically make data safe. Student-level analysis belongs in a separate school-reviewed engagement.")]),
        ("Who decides what", [("table", (["Role", "Responsibility"], [("Andrew", "Demonstrates and teaches the workflow."), ("School leadership", "Selects the tool, account, rules, and escalation contact."), ("Teachers", "Protect information, review results, and make final decisions."), ("Qualified specialists", "Handle privacy, legal, security, special-education, labor, procurement, and regulated questions.")], [1.6, 4.9]))]),
        ("Follow-up reporting", [("p", "Leadership receives cohort-level counts and medians, not participant names, task content, or identifiable comments.")]),
    ]
    save_google_doc(DOCX_DIR / "Leadership-Readiness-Guide.docx", "Teacher Time Back Lab Leadership Readiness Guide", "Complete before the teacher lab is booked", leadership_sections)

    workbook_sections = [
        ("The six steps", [("numbered", ["See the finished product", "Choose one repeated task", "Check the school tool and account", "Build the five fields", "Practice and review", "Save and measure"])]),
        ("Choose one repeated task", [("p", "My task:"), ("p", "How often I do it:"), ("p", "Estimated minutes now:")]),
        ("Check the school tool and account", [("p", "Exact AI product:"), ("p", "Staff account or login path:"), ("p", "School contact if I am unsure:")]),
        ("Build the five fields", [("numbered", [f"{label}:" for _, label, _ in FIVE_FIELDS]), ("p", "Use only the text or files I give you for this task. If the information is not there, mark NEEDS TEACHER INPUT instead of guessing.")]),
        ("Review before use", [("bullets", ["Facts and calculations", "Audience and appropriateness", "Bias and assumptions", "Missing information", "Final wording and instructional judgment"])]),
        ("Two-week measure", [("table", (["Date", "Task", "Before", "Using", "Review", "Result"], [["", "", "", "", "", ""] for _ in range(8)], [0.7, 1.4, 0.8, 0.8, 0.8, 1.0])), ("p", "Decision: Keep, revise, or stop.")]),
    ]
    save_google_doc(DOCX_DIR / "Teacher-Workbook.docx", "Teacher Time Back Lab Workbook", "Build one assistant you can use again", workbook_sections)

    platform_sections = [
        ("Pilot record", [("p", "School or program:"), ("p", "Pilot date:"), ("p", "Exact AI product:"), ("p", "Staff account type and login path:"), ("p", "School host and escalation contact:")]),
        ("Feature decisions", [("table", (["Feature", "Allowed?", "Local note"], [[x, "Yes / No / Ask", ""] for x in ["Paste practice text", "Upload files", "Web browsing or search", "Memory or chat history", "Sharing assistants", "Connectors or external apps", "Saving participant work"]], [2.0, 1.5, 3.0]))]),
        ("Test before the session", [("bullets", ["Participant login works on the school network", "The five-part setup can be saved or copied", "Practice text can be entered without unapproved features", "The facilitator can demonstrate in the same account type", "The school host knows how participants request help"])]),
        ("Minimum first-lab setup", [("p", "The first lab can run without uploads, browsing, connectors, or sharing. Teachers can paste practice text and save instructions inside the school-named account.")]),
    ]
    save_google_doc(DOCX_DIR / "Platform-Setup-Record.docx", "Teacher Time Back Lab Platform Setup Record", "Complete with the school host before the session", platform_sections)


def main():
    pilot_overview_pdf()
    leadership_pdf()
    prep_pdf()
    workbook_pdf()
    practice_pack_pdf()
    safety_card_pdf()
    platform_record_pdf()
    facilitator_pdf()
    build_docx_suite()
    print(f"Built 8 PDFs in {PDF_DIR}")
    print(f"Built 4 Google Docs-ready DOCX files in {DOCX_DIR}")


if __name__ == "__main__":
    main()
